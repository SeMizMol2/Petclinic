from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


WORKSPACE = Path(r"E:\project\Petclinic-fixed\Petclinic")
INPUT_DOCX = WORKSPACE / "temp-report-copy.docx"
OUTPUT_DOCX = WORKSPACE / "ระบบจัดการข้อมูลการรักษาและบริการสัตว์เลี้ยง_ฉบับอัปเดต.docx"


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def replace_paragraph_text(paragraph: Paragraph, text: str) -> None:
    paragraph.text = text


def clear_table_data(table) -> None:
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)


def add_table_row(table, values) -> None:
    row = table.add_row()
    for idx, value in enumerate(values):
        row.cells[idx].text = str(value)


def insert_table_after_paragraph(document: Document, paragraph: Paragraph, rows, cols=6, style=None):
    table = document.add_table(rows=rows, cols=cols)
    if style is not None:
        table.style = style
    paragraph._p.addnext(table._tbl)
    return table


def build_table(table, header, body_rows) -> None:
    clear_table_data(table)
    for idx, text in enumerate(header):
        table.rows[0].cells[idx].text = text
    for row in body_rows:
        add_table_row(table, row)


def find_paragraph(doc: Document, exact_text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"Paragraph not found: {exact_text}")


doc = Document(str(INPUT_DOCX))

# 1) Update scope section for owner features.
replace_paragraph_text(find_paragraph(doc, "1.3.2.5 ได้รับการแจ้งเตือนการนัดหมายผ่านอีเมล"), "1.3.2.5 ดูข้อมูลการนัดหมายของตนเอง")
para_125 = find_paragraph(doc, "1.3.2.5 ดูข้อมูลการนัดหมายของตนเอง")
new_para = insert_paragraph_after(new_para := para_125, "1.3.2.6 ได้รับการแจ้งเตือนการนัดหมายและการเปลี่ยนแปลงสถานะผ่านอีเมล", para_125.style)
replace_paragraph_text(find_paragraph(doc, "1.3.2.6 ดูประวัติการรักษาและรายละเอียดการรับบริการ"), "1.3.2.7 ดูประวัติการรักษาและรายละเอียดการรับบริการ")
replace_paragraph_text(find_paragraph(doc, "1.3.2.7 ดูรายงานสรุปการรักษาและบริการของสัตว์เลี้ยงตนเอง"), "1.3.2.8 ดูรายงานสรุปการรักษาและบริการของสัตว์เลี้ยงตนเอง")

# 2) Add revision notes for DFD / ERD / Data Dictionary.
dfd_heading = find_paragraph(doc, "3.1.3 Data Flow Diagram Level 1")
dfd_note = insert_paragraph_after(
    dfd_heading,
    "หมายเหตุการปรับปรุง: DFD Level 1 ฉบับปรับปรุงให้สอดคล้องกับระบบจริง โดย Process 6 ใช้สำหรับให้ผู้ดูแลระบบหรือเจ้าหน้าที่คลินิกสร้างและจัดการข้อมูลนัดหมายของสัตว์เลี้ยงในระบบ พร้อมรองรับการแจ้งเตือนผลการนัดหมายผ่านอีเมล",
    dfd_heading.style,
)
dfd_note = insert_paragraph_after(
    dfd_note,
    "เพิ่มเติม: Process 7 บันทึกการรักษาและ Process 8 บันทึกประวัติการฉีดวัคซีน อ้างอิงข้อมูลสัตวแพทย์ผ่าน vet_id และเชื่อมโยงไปยังข้อมูลสัตวแพทย์เพื่อใช้แสดงชื่อสัตวแพทย์ในรายงานและหน้าจอแสดงผล",
    dfd_heading.style,
)

erd_heading = find_paragraph(doc, "3.1.4 Entity Relationship Diagram")
erd_note = insert_paragraph_after(
    erd_heading,
    "หมายเหตุการปรับปรุง: ERD ฉบับใช้งานจริงยกเลิกการใช้ owner_address ใน tb_owner และใช้ vet_id เป็น Foreign Key ใน tb_treatment, tb_surgery และ tb_vaccine_rec เพื่อเชื่อมกับ tb_veterinarian",
    erd_heading.style,
)
insert_paragraph_after(
    erd_note,
    "เพิ่มเติม: ฟิลด์ appt_status ของ tb_appointment ใช้งานจริงด้วยค่า รอ, ยืนยัน และ ยกเลิก เพื่อใช้ติดตามสถานะการนัดหมายที่คลินิกบันทึกไว้ในระบบ",
    erd_heading.style,
)

dictionary_heading = find_paragraph(doc, "3.1.5 พจนานุกรมข้อมูล (Data Dictionary)")
dict_note = insert_paragraph_after(
    dictionary_heading,
    "หมายเหตุการปรับปรุง: พจนานุกรมข้อมูลในส่วนนี้ปรับให้ตรงกับฐานข้อมูลจริงของระบบปัจจุบัน ซึ่งใช้งานทั้งหมด 15 ตารางหลัก และเพิ่มตารางรายละเอียดการรักษา (tb_treatment_detail) เพื่อรองรับบริการหลายรายการต่อการรักษา 1 ครั้ง",
    dictionary_heading.style,
)
insert_paragraph_after(
    dict_note,
    "ทั้งนี้ ฟิลด์ owner_address ถูกถอดออกจากการใช้งาน และข้อมูลสัตวแพทย์ในตารางธุรกรรมใช้การอ้างอิงผ่าน vet_id",
    dictionary_heading.style,
)

# 3) Update data dictionary table titles to reflect 15 tables.
replace_paragraph_text(find_paragraph(doc, "ตารางที่ 3-9 ข้อมูลการผ่าตัด (tb_surgery)"), "ตารางที่ 3-10 ข้อมูลการผ่าตัด (tb_surgery)")
replace_paragraph_text(find_paragraph(doc, "ตารางที่ 3-10 ประวัติการฉีดวัคซีน (tb_vaccine_rec)"), "ตารางที่ 3-11 ประวัติการฉีดวัคซีน (tb_vaccine_rec)")
replace_paragraph_text(find_paragraph(doc, "ตารางที่ 3-11 ข้อมูลใบเสร็จ/ใบแจ้งหนี้ (tb_receipt)"), "ตารางที่ 3-12 ข้อมูลใบเสร็จ/ใบแจ้งหนี้ (tb_receipt)")
replace_paragraph_text(find_paragraph(doc, "ตารางที่ 3-12 รายละเอียดใบเสร็จ (tb_receipt_detail)"), "ตารางที่ 3-13 รายละเอียดใบเสร็จ (tb_receipt_detail)")
replace_paragraph_text(find_paragraph(doc, "ตารางที่ 3-13 ข้อมูลรายจ่าย (tb_expense)"), "ตารางที่ 3-14 ข้อมูลรายจ่าย (tb_expense)")
replace_paragraph_text(find_paragraph(doc, "ตารางที่ 3-14 ข้อมูลสัตวแพทย์ (tb_veterinarian)"), "ตารางที่ 3-15 ข้อมูลสัตวแพทย์ (tb_veterinarian)")

# 4) Update owner table.
owner_table = doc.tables[11]
build_table(
    owner_table,
    ["No", "FieldName", "Description", "Data Type", "Size", "Key"],
    [
        [1, "owner_id", "รหัสเจ้าของ", "VARCHAR", "20", "PK"],
        [2, "owner_name", "ชื่อ-นามสกุล", "VARCHAR", "100", ""],
        [3, "owner_email", "อีเมลเจ้าของสัตว์เลี้ยง", "VARCHAR", "255", ""],
        [4, "owner_tel", "เบอร์โทรศัพท์", "VARCHAR", "15", ""],
        [5, "user_id", "รหัสผู้ใช้งาน", "VARCHAR", "10", "FK"],
        [6, "profile_pic", "รูปโปรไฟล์เจ้าของสัตว์เลี้ยง", "TEXT", "", ""],
        [7, "create_datetime", "วันเวลาที่สร้างข้อมูล", "DATETIME", "", ""],
        [8, "update_datetime", "วันเวลาที่แก้ไขข้อมูล", "DATETIME", "", ""],
    ],
)

# 5) Update appointment table.
appointment_table = doc.tables[13]
build_table(
    appointment_table,
    ["No", "FieldName", "Description", "Data Type", "Size", "Key"],
    [
        [1, "appt_id", "รหัสนัดหมาย", "VARCHAR", "10", "PK"],
        [2, "appt_date", "วันที่นัด", "DATE", "", ""],
        [3, "appt_time", "เวลานัด", "TIME", "", ""],
        [4, "appt_reason", "เหตุผลหรืออาการเบื้องต้น", "TEXT", "", ""],
        [5, "cancel_reason", "เหตุผลการยกเลิกนัดหมาย", "TEXT", "", ""],
        [6, "appt_status", "สถานะนัดหมาย (รอ, ยืนยัน, ยกเลิก)", "VARCHAR", "20", ""],
        [7, "pet_id", "รหัสสัตว์เลี้ยง", "VARCHAR", "10", "FK"],
        [8, "create_datetime", "เวลาที่สร้างข้อมูล", "DATETIME", "", ""],
        [9, "update_datetime", "เวลาที่แก้ไขข้อมูล", "DATETIME", "", ""],
    ],
)

# 6) Update treatment table to actual program structure.
treatment_table = doc.tables[14]
build_table(
    treatment_table,
    ["No", "FieldName", "Description", "Data Type", "Size", "Key"],
    [
        [1, "treatment_id", "รหัสการรักษา", "VARCHAR", "15", "PK"],
        [2, "pet_id", "รหัสสัตว์เลี้ยง", "VARCHAR", "20", "FK"],
        [3, "user_id", "รหัสผู้ใช้งานที่บันทึกรายการ", "VARCHAR", "20", "FK"],
        [4, "vet_id", "รหัสสัตวแพทย์", "VARCHAR", "10", "FK"],
        [5, "symptom", "อาการสำคัญ", "TEXT", "", ""],
        [6, "diagnosis", "ผลการวินิจฉัย", "TEXT", "", ""],
        [7, "treatment_date", "วันที่เข้ารักษา", "DATETIME", "", ""],
        [8, "total_amount", "ยอดรวมค่ารักษา", "NUMERIC", "10,2", ""],
    ],
)

# 7) Insert treatment detail title and table after treatment title.
treatment_title = find_paragraph(doc, "ตารางที่ 3-8 ประวัติการรักษา (tb_treatment)")
treatment_detail_title = insert_paragraph_after(treatment_title, "ตารางที่ 3-9 รายละเอียดการรักษา (tb_treatment_detail)", treatment_title.style)
new_table = insert_table_after_paragraph(doc, treatment_detail_title, 1, 6, style=doc.tables[14].style)
build_table(
    new_table,
    ["No", "FieldName", "Description", "Data Type", "Size", "Key"],
    [
        [1, "detail_id", "รหัสรายละเอียด", "INTEGER", "", "PK"],
        [2, "treatment_id", "รหัสการรักษา", "VARCHAR", "15", "FK"],
        [3, "service_id", "รหัสบริการ", "VARCHAR", "10", "FK"],
        [4, "quantity", "จำนวนรายการ", "INTEGER", "", ""],
        [5, "price", "ราคาต่อหน่วย", "NUMERIC", "10,2", ""],
    ],
)

# 8) Update Home / user / admin appointment UI descriptions.
replace_paragraph_text(
    find_paragraph(doc, "จากภาพประกอบที่ 3-20 คือหน้าแรกเมื่อเข้าสู่เว็บไซต์จะแสดงข้อมูลเบื้องต้นเกี่ยวกับการให้บริการของคลินิก ข้อมูลติดต่อ ที่อยู่ เวลาที่ทำการ และ มีปุ่มสำหรับสมัครสมาชิกและเข้าสู่ระบบ"),
    "จากภาพประกอบที่ 3-20 คือหน้าแรกเมื่อเข้าสู่เว็บไซต์ ซึ่งจะแสดงข้อมูลจริงของคลินิกจากฐานข้อมูล เช่น ข้อมูลติดต่อ ที่อยู่ เวลาที่ทำการ และรายการบริการพร้อมค่าบริการเบื้องต้น รวมถึงมีปุ่มสำหรับสมัครสมาชิกและเข้าสู่ระบบ",
)
replace_paragraph_text(
    find_paragraph(doc, "จากภาพประกอบที่ 3-23 คือหน้าที่จะแสดงข้อมูลผู้ใช้งานหลังจากที่เข้าสู่ระบบโดยจะมีปุ่มจัดการข้อมูลสัตว์เลี้ยงที่จะเชื่อมไปยังหน้าที่จะสามารถเพิ่ม/ลบ/แก้ไข ข้อมูลสัตว์เลี้ยงได้ และปุ่มโปรไฟล์เจ้าของสัตว์เลี้ยงที่เมื่อกดแล้วผู้ใช้สามารถแก้ไขข้อมูลส่วนตัวได้"),
    "จากภาพประกอบที่ 3-23 คือหน้าแรกหลังเข้าสู่ระบบของผู้ใช้งาน ซึ่งทำหน้าที่เป็นศูนย์กลางสำหรับเข้าถึงข้อมูลส่วนตัว ข้อมูลสัตว์เลี้ยง การนัดหมาย ประวัติการรักษา และประวัติการชำระเงินของเจ้าของสัตว์เลี้ยงในบัญชีเดียว",
)
replace_paragraph_text(
    find_paragraph(doc, "จากภาพประกอบที่ 3-24 คือหน้าที่ไว้สำหรับแก้ไขข้อมูลส่วนตัวในส่วนของผู้ใช้งาน"),
    "จากภาพประกอบที่ 3-24 คือหน้าที่สำหรับแก้ไขข้อมูลส่วนตัวของผู้ใช้งาน เช่น ชื่อ อีเมล เบอร์โทรศัพท์ และรูปโปรไฟล์ เพื่อให้ข้อมูลติดต่อสำหรับการแจ้งเตือนการนัดหมายเป็นปัจจุบัน",
)
replace_paragraph_text(
    find_paragraph(doc, "จากภาพประกอบที่ 3-32 คือหน้าที่ไว้สำหรับแพทย์หรือผู้ดูแลระบบจัดการนัดหมายวันเข้ารับการรักษาของสัตว์เลี้ยง"),
    "จากภาพประกอบที่ 3-32 คือหน้าที่สำหรับแพทย์หรือผู้ดูแลระบบจัดการข้อมูลนัดหมายของสัตว์เลี้ยง สามารถสร้างนัดหมายใหม่ แก้ไขข้อมูลนัดหมาย และเปลี่ยนสถานะเป็น รอ, ยืนยัน หรือ ยกเลิก พร้อมส่งอีเมลแจ้งผลการนัดหมายถึงเจ้าของสัตว์เลี้ยงได้",
)

# 9) Add missing user-facing UI sections before 3.2.10.
add_pet_heading = find_paragraph(doc, "3.2.10 หน้าประวัติการรักษาและบริการ")
new_ui = add_pet_heading.insert_paragraph_before("3.2.7 หน้ารายการสัตว์เลี้ยงของผู้ใช้งาน")
new_ui.style = add_pet_heading.style
new_ui = insert_paragraph_after(
    new_ui,
    "หน้ารายการสัตว์เลี้ยงของผู้ใช้งานใช้สำหรับแสดงสัตว์เลี้ยงทั้งหมดในบัญชีเดียวกัน โดยผู้ใช้สามารถเลือกดูรายละเอียด แก้ไขข้อมูล และเชื่อมต่อไปยังประวัติการรักษาหรือข้อมูลนัดหมายของสัตว์เลี้ยงแต่ละตัวได้",
    doc.paragraphs[0].style,
)
new_ui = insert_paragraph_after(new_ui, "3.2.8 หน้าการนัดหมายของผู้ใช้งาน", add_pet_heading.style)
new_ui = insert_paragraph_after(
    new_ui,
    "หน้าการนัดหมายของผู้ใช้งานใช้สำหรับตรวจสอบข้อมูลนัดหมายที่คลินิกบันทึกไว้ให้กับสัตว์เลี้ยงของตนเอง พร้อมติดตามสถานะนัดหมายในรูปแบบ รอ, ยืนยัน และ ยกเลิก",
    doc.paragraphs[0].style,
)
new_ui = insert_paragraph_after(new_ui, "3.2.9 หน้าประวัติการชำระเงินของผู้ใช้งาน", add_pet_heading.style)
insert_paragraph_after(
    new_ui,
    "หน้าประวัติการชำระเงินของผู้ใช้งานใช้สำหรับตรวจสอบใบเสร็จ ยอดชำระ สถานะการชำระเงิน และรายละเอียดรายการที่เกี่ยวข้องกับการรักษาและบริการของสัตว์เลี้ยงแต่ละตัว",
    doc.paragraphs[0].style,
)

doc.save(str(OUTPUT_DOCX))
print(OUTPUT_DOCX)
